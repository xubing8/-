import sys
import sqlite3
import requests
import hashlib
import random
import os
import pandas as pd
from PyQt5 import QtWidgets, QtCore
from transformers import BertTokenizer, BertModel
import torch
import faiss
import numpy as np
import pickle
import jieba
from langdetect import detect, LangDetectException
import time
from requests.exceptions import RequestException
from LAC import LAC  # 导入 LAC
import re
from docx import Document
from docx.shared import Inches
import warnings
import logging
import sys
import subprocess
import pkg_resources
import requests
import logging
from tenacity import retry, wait_exponential, stop_after_attempt, retry_if_exception_type
from requests.exceptions import HTTPError

import subprocess
import sys

# 自定义异常，用于捕获速率限制错误
class RateLimitError(Exception):
    pass
# 定义所需的库
required_packages = {'openai==0.28','sentence-splitter','ratelimit','tenacity','anthropic', 'python-dotenv'}

# 获取已安装的库
installed_packages = {pkg.key for pkg in pkg_resources.working_set}

# 找出缺失的库
missing_packages = required_packages - installed_packages

if missing_packages:
    print(f"缺失的库: {missing_packages}，正在安装...")
    try:
        # 使用当前 Python 解释器的 pip 安装缺失的库
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', *missing_packages])
        print("安装完成。")
    except subprocess.CalledProcessError as e:
        print(f"安装失败: {e}")
        sys.exit(1)

# 现在可以安全地导入所需的库
import anthropic
from dotenv import load_dotenv
from ratelimit import limits, sleep_and_retry
from tenacity import retry, wait_exponential, stop_after_attempt, retry_if_exception_type
from sentence_splitter import split_text_into_sentences
import openai

load_dotenv()

# 配置日志记录
logging.basicConfig(
    filename='translator_app.log',
    filemode='a',
    format='%(asctime)s - %(levelname)s - %(message)s',
    level=logging.DEBUG
)

os.environ['KMP_DUPLICATE_LIB_OK'] = 'TRUE'

warnings.filterwarnings('ignore')

import nltk
class TranslatorApp(QtWidgets.QWidget):


    def __init__(self):
        super().__init__()

        # 获取百度翻译API的APP_ID和SECRET_KEY
        # 获取百度翻译API的APP_ID和SECRET_KEY
        self.APP_ID = os.getenv('BAIDU_APP_ID', '20231208001905106')  # 请确保设置环境变量或替换为您的APPID
        self.SECRET_KEY = os.getenv('BAIDU_SECRET_KEY', 'pJbqmZ32GaoM9nBQ8q_0')  # 请确保设置环境变量或替换为您的密钥
        self.term_translation_cache = {}  # 添加术语翻译缓存
        self.messages = []  # 初始化 messages 列表
        self.init_ui()
        self.init_db()
        self.init_lac()  # 初始化 LAC
        self.load_terminology_into_jieba()  # 加载术语库到 Jieba
        self.init_model()
        self.init_faiss()
        self.init_claude_api()  
        self.init_kimi_api()  # 初始化 Kimi API
        self.init_openai_api()  # 初始化 OpenAI API
        self.init_nltk()  # 初始化 NLTK 资源
        self.initialize_kimi_messages()  # 初始化 Kimi API 的消息列表
        self.extract_and_store_key_terms()  # 提取并存储关键名词
        self.show()
    
    def init_nltk(self):
        try:
            nltk.download('punkt', quiet=True)
            logging.info("成功下载并初始化 NLTK 'punkt' 资源。")
        except Exception as e:
            logging.error(f"初始化 NLTK 资源时出错: {e}")

    def initialize_kimi_messages(self):
        """
        初始化 Kimi API 的消息列表，添加系统消息以设定上下文。
        """
        self.messages = [
            {
                "role": "system",
                "content": (
                    "你是一个专业的翻译助手，擅长将中文翻译为多种语言，保持上下文的连贯性和准确性。"
                    "请确保翻译结果自然流畅，符合目标语言的语法和表达习惯。"
                )
            }
        ]

    
    def init_kimi_api(self):
        """
        初始化 Kimi API 客户端
        """
        try:
            self.kimi_api_key = os.getenv('KIMI_API_KEY')
            self.kimi_base_url = os.getenv('KIMI_BASE_URL', 'https://api.moonshot.cn/v1')
            if not self.kimi_api_key:
                raise ValueError("KIMI_API_KEY 环境变量未设置。")
            logging.info("成功初始化 Kimi API 客户端。")
        except Exception as e:
            logging.error(f"初始化 Kimi API 客户端失败: {e}")
            self.kimi_api_key = None

    def init_ui(self):
        self.setWindowTitle('翻译记忆辅助工具')

        # 输入框和标签
        self.input_label = QtWidgets.QLabel('请输入文本：')
        self.input_text = QtWidgets.QTextEdit()

        self.output_label = QtWidgets.QLabel('翻译结果：')
        self.output_text = QtWidgets.QTextEdit()
        self.output_text.setReadOnly(True)

        # 语言选择
        self.from_label = QtWidgets.QLabel('源语言：')
        self.from_lang = QtWidgets.QComboBox()
        # 添加印地语和阿拉伯语
        self.from_lang.addItems(['自动检测', '中文', '英文', '日文', '韩文', '法文', '德文', '俄文', '西班牙文', '印地语', '阿拉伯语'])
        self.from_lang.setCurrentIndex(0)  # 默认自动检测

        self.to_label = QtWidgets.QLabel('目标语言：')
        self.to_lang = QtWidgets.QComboBox()
        # 添加印地语和阿拉伯语
        self.to_lang.addItems(['中文', '英文', '日文', '韩文', '法文', '德文', '俄文', '西班牙文', '印地语', '阿拉伯语'])
        self.to_lang.setCurrentIndex(1)  # 默认英文

        # 翻译引擎选择
        self.engine_label = QtWidgets.QLabel('翻译引擎：')
        self.engine_select = QtWidgets.QComboBox()
        self.engine_select.addItems(['百度翻译API', 'Kimi API','OpenAI API'])  # 添加 'Kimi API'
        self.engine_select.setCurrentIndex(1)  # 默认选择 Kimi API

        # 按钮
        self.translate_button = QtWidgets.QPushButton('翻译')
        self.translate_button.clicked.connect(self.translate_text)

        self.import_button = QtWidgets.QPushButton('导入翻译记忆库')
        self.import_button.clicked.connect(self.import_memory)

        self.export_button = QtWidgets.QPushButton('导出翻译记忆库')
        self.export_button.clicked.connect(self.export_memory)

        self.view_memory_button = QtWidgets.QPushButton('查看记忆库')
        self.view_memory_button.clicked.connect(self.view_memory)

        self.import_word_button = QtWidgets.QPushButton('导入 Word 文档')
        self.import_word_button.clicked.connect(self.import_word_document)

        self.export_word_button = QtWidgets.QPushButton('导出为 Word')
        self.export_word_button.clicked.connect(self.export_to_word)

        # 新增清空记忆库按钮
        self.clear_memory_button = QtWidgets.QPushButton('清空记忆库')
        self.clear_memory_button.clicked.connect(self.clear_memory)

        # 布局
        layout = QtWidgets.QGridLayout()
        layout.addWidget(self.input_label, 0, 0)
        layout.addWidget(self.input_text, 1, 0, 1, 2)
        layout.addWidget(self.from_label, 2, 0)
        layout.addWidget(self.from_lang, 2, 1)
        layout.addWidget(self.to_label, 3, 0)
        layout.addWidget(self.to_lang, 3, 1)
        layout.addWidget(self.engine_label, 4, 0)
        layout.addWidget(self.engine_select, 4, 1)
        layout.addWidget(self.translate_button, 5, 0)
        layout.addWidget(self.import_button, 5, 1)
        layout.addWidget(self.import_word_button, 6, 0)
        layout.addWidget(self.export_word_button, 6, 1)
        layout.addWidget(self.export_button, 7, 0)
        layout.addWidget(self.view_memory_button, 7, 1)
        layout.addWidget(self.clear_memory_button, 8, 0, 1, 2)  # 清空记忆库按钮占据两列
        layout.addWidget(self.output_label, 9, 0)
        layout.addWidget(self.output_text, 10, 0, 1, 2)

        self.setLayout(layout)
        self.resize(800, 600)

    def init_openai_api(self):
        """
        初始化 OpenAI API 客户端
        """
        try:
            self.openai_api_key = 'sk-vdUAyw8Wzflt60b23102A35dC98241C48429Df2c5e825aDb'  # 直接设置 API Key
            self.openai_base_url = 'https://api.wlai.vip/v1'  # 设置为您的中转接口地址

            if not self.openai_api_key:
                # 提示用户输入 API Key
                self.openai_api_key, ok = QtWidgets.QInputDialog.getText(
                    self, 'OpenAI API Key', '请输入 OpenAI API Key：'
                )
                if not ok or not self.openai_api_key:
                    raise ValueError("OpenAI API Key 未提供。")

            openai.api_key = self.openai_api_key
            openai.api_base = self.openai_base_url
            logging.info("成功初始化 OpenAI API 客户端。")
        except Exception as e:
            logging.error(f"初始化 OpenAI API 客户端失败: {e}")
            self.openai_api_key = None


    def init_claude_api(self):
        """
        初始化Anthropic Claude API客户端
        """
        try:
            self.claude_client = anthropic.Anthropic(
                api_key=os.getenv('ANTHROPIC_API_KEY'),
                base_url=os.getenv('ANTHROPIC_BASE_URL', 'https://cc.plusai.me')
            )
            logging.info("成功初始化Claude API客户端。")
        except Exception as e:
            logging.error(f"初始化Claude API客户端失败: {e}")
            self.claude_client = None

    def claude_translate(self, query, from_lang, to_lang, model="claude-3-5-sonnet"):
        """
        使用非官方API进行翻译
        """
        base_url = 'https://cc.plusai.me/v1/chat/completions'  # 请根据实际API端点进行修改
        api_key = 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJleHAiOjE3MjY5MzI5OTksInN1YiI6ImYxNmQ5NjY4ZjFmNDQwZDBhN2I2ZTM3NGQ5NTJkZTk1IiwiaWF0IjoxNzI2OTA0MTk5LCJ0b2tlbl90eXBlIjoiYWNjZXNzX3Rva2VuIiwic3Vic2NyaXB0aW9uX2lkIjoiNjZlYmQ2ZWMzOTFkMjkyMDhjNTQ0OWFmIiwiYWNjb3VudF9pZCI6IjY1ZWRiY2E3OTk4NWIzZDZhMzZiYTVjYiJ9.qH-a2Vq0mFLA2gprDoCEBMjJYFPkJQJSGIfhD3aLc0E'

        headers = {
            'Content-Type': 'application/json',
            'x-api-key': api_key
        }

        lang_map = {
            '自动检测': 'auto',
            '中文': 'Chinese',
            '英文': 'English',
            '日文': 'Japanese',
            '韩文': 'Korean',
            '法文': 'French',
            '德文': 'German',
            '俄文': 'Russian',
            '西班牙文': 'Spanish',
            # 根据Claude API支持的语言进行扩展
        }

        prompt = (
            f"请将以下文本从{lang_map.get(from_lang, from_lang)}翻译为{lang_map.get(to_lang, to_lang)}，"
            "只需给出翻译结果，不要添加任何解释或注释。\n\n"
            f"源文本：\n{query}\n\n译文："
        )

        data = {
            "model": model,
            "prompt": prompt,
            "max_tokens": 1024,
            "temperature": 0.3
        }

        try:
            logging.debug(f"发送到Claude API的请求: URL={base_url}, Headers={headers}, Data={data}")
            response = requests.post(base_url, headers=headers, json=data)

            if response.status_code == 200:
                result = response.json()
                translation = result.get('text', '').strip()
                logging.info(f"Claude 翻译结果: {translation}")
                return translation
            else:
                error_detail = response.json().get('detail', '未知错误')
                logging.error(f"调用Claude API出错: Error code: {response.status_code} - {response.json()}")
                return f'翻译失败：{error_detail}'
        except Exception as e:
            logging.error(f"调用Claude API出错: {e}")
            return f'翻译失败：{str(e)}'



    def clear_memory(self):
        reply = QtWidgets.QMessageBox.question(
            self, '确认清空',
            '您确定要清空所有翻译记忆库吗？此操作无法撤销。',
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No,
            QtWidgets.QMessageBox.No
        )
        if reply == QtWidgets.QMessageBox.Yes:
            try:
                # 清空句子翻译表
                self.cursor.execute("DELETE FROM sentence_translations")
                # 清空词汇翻译表
                self.cursor.execute("DELETE FROM word_translations")
                # 清空术语库表
                self.cursor.execute("DELETE FROM key_terms")
                self.conn.commit()

                # 重新初始化FAISS索引和ID映射
                self.index.reset()
                self.id_map = {}
                faiss.write_index(self.index, self.faiss_index_file)
                with open(self.id_map_file, 'wb') as f:
                    pickle.dump(self.id_map, f)

                # 重新加载 Jieba 词典
                self.load_terminology_into_jieba()

                QtWidgets.QMessageBox.information(self, '成功', '翻译记忆库已清空。')
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, '错误', f'清空记忆库失败：{str(e)}')

    def init_db(self):
        self.conn = sqlite3.connect('translation_memory.db')
        self.cursor = self.conn.cursor()
        # 创建句子翻译表
        self.cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='sentence_translations'")
        if not self.cursor.fetchone():
            self.cursor.execute('''
                CREATE TABLE sentence_translations (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    source TEXT NOT NULL,
                    target TEXT NOT NULL,
                    from_lang TEXT NOT NULL,
                    to_lang TEXT NOT NULL,
                    source_vector BLOB
                )
            ''')
            self.conn.commit()

        # 创建词汇翻译表
        self.cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='word_translations'")
        if not self.cursor.fetchone():
            self.cursor.execute('''
                CREATE TABLE word_translations (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    source TEXT NOT NULL,
                    target TEXT NOT NULL,
                    from_lang TEXT NOT NULL,
                    to_lang TEXT NOT NULL
                )
            ''')
            self.conn.commit()

        # 创建术语库表
        self.cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='key_terms'")
        if not self.cursor.fetchone():
            self.cursor.execute('''
                CREATE TABLE key_terms (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    source_term TEXT NOT NULL UNIQUE,
                    target_term TEXT NOT NULL,
                    from_lang TEXT NOT NULL,
                    to_lang TEXT NOT NULL
                )
            ''')
            self.conn.commit()

    def init_lac(self):
        try:
            self.lac = LAC(mode='lac')  # 使用默认模式，包含词性标注和命名实体识别
        except Exception as e:
            print(f"LAC 初始化失败: {e}")
            self.lac = None

    def load_terminology_into_jieba(self):
        """
        从 key_terms 表中获取所有源术语为中文的词汇，并添加到 Jieba 词典中。
        """
        self.cursor.execute("SELECT source_term FROM key_terms WHERE from_lang='中文'")
        words = [row[0] for row in self.cursor.fetchall()]
        if words:
            # 将词汇写入用户词典文件
            with open('user_dict.txt', 'w', encoding='utf-8') as f:
                for word in words:
                    f.write(f"{word}\n")
            jieba.load_userdict('user_dict.txt')

    def openai_translate(self, query, from_lang, to_lang, model="gpt-4o"):
        """
        使用 OpenAI 的 ChatCompletion API 进行翻译
        """
        if not self.openai_api_key:
            logging.error("OpenAI API Key 未设置。")
            return '翻译失败：OpenAI API Key 未设置。'

        lang_map = {
            '自动检测': 'auto',
            '中文': 'Chinese',
            '英文': 'English',
            '日文': 'Japanese',
            '韩文': 'Korean',
            '法文': 'French',
            '德文': 'German',
            '俄文': 'Russian',
            '西班牙文': 'Spanish',
            '印地语': 'Hindi',
            '阿拉伯语': 'Arabic',
        }

        prompt = (
            f"Please translate the following text from {lang_map.get(from_lang, from_lang)} to {lang_map.get(to_lang, to_lang)}. "
            "Provide the translation only, without any explanations or comments.\n\n"
            f"Text:\n{query}\n\nTranslation:"
        )

        try:
            response = openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=1024
            )
            translation = response.choices[0].message.content.strip()
            logging.info(f"OpenAI 翻译结果: {translation}")
            return translation
        except Exception as e:
            logging.error(f"调用 OpenAI API 出错: {e}")
            return f'翻译失败：{str(e)}'

    def init_model(self):
        model_name = 'bert-base-chinese'  # 根据需要更换模型
        self.tokenizer = BertTokenizer.from_pretrained(model_name)
        self.model = BertModel.from_pretrained(model_name)
        self.model.eval()
        # 如果有GPU可用，使用GPU
        self.device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        self.model.to(self.device)

    def init_faiss(self):
        # 初始化FAISS索引
        self.faiss_index_file = 'faiss_index.bin'
        self.id_map_file = 'id_map.pkl'
        if os.path.exists(self.faiss_index_file) and os.path.exists(self.id_map_file):
            # 加载FAISS索引和ID映射
            self.index = faiss.read_index(self.faiss_index_file)
            with open(self.id_map_file, 'rb') as f:
                self.id_map = pickle.load(f)
        else:
            # 创建新的FAISS索引
            self.index = faiss.IndexFlatL2(768)  # BERT的隐藏层维度是768
            self.id_map = {}
            self.build_faiss_index()

    def build_faiss_index(self):
        # 从数据库中加载所有向量并添加到FAISS索引
        self.cursor.execute("SELECT id, source_vector FROM sentence_translations WHERE source_vector IS NOT NULL")
        rows = self.cursor.fetchall()
        vectors = []
        ids = []
        for row in rows:
            id_, vector_blob = row
            if vector_blob:
                vector = pickle.loads(vector_blob)
                vectors.append(vector)
                ids.append(id_)
        if vectors:
            vectors_np = np.array(vectors).astype('float32')
            self.index.add(vectors_np)
            for idx, id_ in enumerate(ids):
                self.id_map[idx] = id_
            # 保存ID映射
            with open(self.id_map_file, 'wb') as f:
                pickle.dump(self.id_map, f)
            # 保存FAISS索引
            faiss.write_index(self.index, self.faiss_index_file)

    def extract_and_store_key_terms(self):
        if not self.lac:
            print("LAC 未初始化，无法提取关键名词。")
            return

        self.cursor.execute("SELECT source, target FROM sentence_translations")
        rows = self.cursor.fetchall()

        for source_text, target_text in rows:
            self.extract_and_store_key_terms_from_sentence(source_text, target_text)

    def extract_and_store_key_terms_from_sentence(self, source_text, target_text):
        if not self.lac:
            print("LAC 未初始化，无法提取关键名词。")
            return

        lac_result = self.lac.run(source_text)
        words, tags = lac_result

        for word, tag in zip(words, tags):
            if tag in ['PER', 'ORG', 'LOC', 'MISC']:  # 人名、组织名、地名等
                # 检查是否已存在
                self.cursor.execute("SELECT id FROM key_terms WHERE source_term=?", (word,))
                if not self.cursor.fetchone():
                    # 获取术语的英文翻译
                    translated_word = self.get_translation_of_term(word)
                    if translated_word:
                        self.cursor.execute('''
                            INSERT INTO key_terms (source_term, target_term, from_lang, to_lang)
                            VALUES (?, ?, ?, ?)
                        ''', (word, translated_word, '中文', '英文'))
        self.conn.commit()
        # 重新加载 Jieba 词典
        self.load_terminology_into_jieba()


    def get_translation_of_term(self, term):
        """
        翻译中文术语为英文。
        使用缓存以减少API调用次数。
        """
        if term in self.term_translation_cache:
            return self.term_translation_cache[term]
        
        try:
            translated_term = self.baidu_translate(term, '中文', '英文')
            # 如果翻译失败，返回原术语
            if not translated_term.startswith('翻译失败') and not translated_term.startswith('请求翻译服务出错'):
                self.term_translation_cache[term] = translated_term
                return translated_term
            else:
                print(f"翻译术语失败: {translated_term}")
                return term  # 回退到原术语
        except Exception as e:
            print(f"获取术语翻译时出错: {e}")
            return term  # 回退到原术语


    def add_sentence_translation(self, source_text, target_text, from_lang, to_lang):
        # 计算源文本的向量
        source_vector = self.compute_vector(source_text)
        if source_vector is not None:
            source_vector_blob = pickle.dumps(source_vector)
        else:
            source_vector_blob = None
        self.cursor.execute('''
            INSERT INTO sentence_translations (source, target, from_lang, to_lang, source_vector)
            VALUES (?, ?, ?, ?, ?)
        ''', (source_text, target_text, from_lang, to_lang, source_vector_blob))
        self.conn.commit()
        inserted_id = self.cursor.lastrowid
        if source_vector is not None:
            # 添加到FAISS索引
            vector_np = np.array([source_vector]).astype('float32')
            self.index.add(vector_np)
            self.id_map[len(self.id_map)] = inserted_id
            # 保存ID映射和FAISS索引
            with open(self.id_map_file, 'wb') as f:
                pickle.dump(self.id_map, f)
            faiss.write_index(self.index, self.faiss_index_file)

        # 提取并存储关键名词
        self.extract_and_store_key_terms_from_sentence(source_text, target_text)

    def add_word_translation(self, source_word, target_word, from_lang, to_lang):
        self.cursor.execute('''
            INSERT INTO word_translations (source, target, from_lang, to_lang)
            VALUES (?, ?, ?, ?)
        ''', (source_word, target_word, from_lang, to_lang))
        self.conn.commit()
        # 如果添加的词汇是中文的，加载到 Jieba
        if from_lang == '中文':
            jieba.add_word(source_word)
            # 也可以更新 user_dict.txt
            with open('user_dict.txt', 'a', encoding='utf-8') as f:
                f.write(f"{source_word}\n")

    def fetch_similar_sentences(self, source_text, top_k=5):
        # 计算源文本的向量
        vector = self.compute_vector(source_text)
        if vector is None:
            return []
        vector_np = np.array([vector]).astype('float32')
        D, I = self.index.search(vector_np, top_k)
        similar_translations = []
        for idx, distance in zip(I[0], D[0]):
            if idx == -1:
                continue
            if idx in self.id_map:
                id_ = self.id_map[idx]
                self.cursor.execute('SELECT source, target FROM sentence_translations WHERE id=?', (id_,))
                row = self.cursor.fetchone()
                if row:
                    similar_translations.append((row[0], row[1], distance))
        return similar_translations

    def compute_vector(self, text):
        try:
            inputs = self.tokenizer(text, return_tensors='pt', truncation=True, max_length=512)
            inputs = {k: v.to(self.device) for k, v in inputs.items()}
            with torch.no_grad():
                outputs = self.model(**inputs)
            # 使用最后一层的平均池化作为句子向量
            vector = outputs.last_hidden_state.mean(dim=1).squeeze().cpu().numpy()
            return vector
        except Exception as e:
            print(f"计算向量时出错: {e}")
            return None

    def find_similar_sentence(self, source_text, top_k=5, threshold=0.8):
        similar_translations = self.fetch_similar_sentences(source_text, top_k)
        best_translation = None
        best_similarity = 0.0
        for source, target, distance in similar_translations:
            # FAISS基于L2距离，简单转换为相似度
            similarity = 1 / (1 + distance)
            if similarity > best_similarity:
                best_similarity = similarity
                best_translation = target
        if best_similarity > threshold:
            return best_translation, best_similarity
        else:
            return None, None

    def replace_words(self, text, to_lang, from_lang):
        if to_lang == '中文':
            # 对中文进行分词
            words = list(jieba.cut(text))
        else:
            # 对英文按空格分词
            words = text.split()

        translated_words = []
        for word in words:
            # 查找词汇翻译
            self.cursor.execute('''
                SELECT target FROM word_translations
                WHERE source=? AND from_lang=? AND to_lang=?
            ''', (word, from_lang, to_lang))
            result = self.cursor.fetchone()
            if result:
                translated_word = result[0]
            else:
                translated_word = word
            translated_words.append(translated_word)

        if to_lang == '中文':
            return ''.join(translated_words)
        else:
            return ' '.join(translated_words)
        
    def baidu_translate(self, query, from_lang, to_lang):
        appid = self.APP_ID
        secretKey = self.SECRET_KEY
        salt = str(random.randint(32768, 65536))

        # Step 1: 拼接字符串 appid + query + salt + secretKey
        sign_str = appid + query + salt + secretKey

        # Step 2: 生成 MD5 签名
        sign = hashlib.md5(sign_str.encode('utf-8')).hexdigest()

        url = 'https://fanyi-api.baidu.com/api/trans/vip/translate'

        lang_map = {
            '自动检测': 'auto',
            '中文': 'zh',
            '英文': 'en'
        }

        params = {
            'q': query,
            'from': lang_map.get(from_lang, 'auto'),
            'to': lang_map.get(to_lang, 'en'),
            'appid': appid,
            'salt': salt,
            'sign': sign
        }
        try:
            response = requests.get(url, params=params)
            time.sleep(0.6)
            result = response.json()

            if 'trans_result' in result:
                return result['trans_result'][0]['dst']
            else:
                return '翻译失败：' + result.get('error_msg', '未知错误')

        except Exception as e:
            return '请求翻译服务出错：' + str(e)

    @retry(wait=wait_exponential(multiplier=1, min=4, max=10), stop=stop_after_attempt(5), retry=retry_if_exception_type(RateLimitError))
    def kimi_translate(self, query, from_lang, to_lang, model="moonshot-v1-32k"):
        """
        使用 Kimi API 进行翻译，并实现多轮对话和重试机制
        """
        if not self.kimi_api_key:
            logging.error("Kimi API Key 未设置。")
            return '翻译失败：Kimi API Key 未设置。'
        
        url = f"{self.kimi_base_url}/chat/completions"
        
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f"Bearer {self.kimi_api_key}"
        }
        
        lang_map = {
            '自动检测': 'auto',
            '中文': 'zh',
            '英文': 'en',
            '日文': 'ja',
            '韩文': 'ko',
            '法文': 'fr',
            '德文': 'de',
            '俄文': 'ru',
            '西班牙文': 'es',
            '印地语': 'hi',
            '阿拉伯语': 'ar',
        }
        
        # 构建用户消息，包含翻译指令和源文本
        user_message = (
            f"请将以下文本从{lang_map.get(from_lang, from_lang)}翻译为{lang_map.get(to_lang, to_lang)}。"
            "保持上下文的一致性和准确性，使用专业术语时请确保其正确性。\n\n"
            f"源文本:\n{query}\n\n译文:"
        )
        
        # 构建消息列表，包括历史对话和当前的用户消息
        messages = self.messages.copy()
        messages.append({
            "role": "user",
            "content": user_message
        })
        
        data = {
            "model": model,
            "messages": messages,
            "temperature": 0.3,
        }
        
        try:
            logging.debug(f"发送到 Kimi API 的请求: URL={url}, Headers={headers}, Data={data}")
            response = requests.post(url, headers=headers, json=data)
            
            # 处理速率限制错误
            if response.status_code == 429:
                raise RateLimitError("Too Many Requests")
            response.raise_for_status()
            
            result = response.json()
            translation = result.get('choices', [{}])[0].get('message', {}).get('content', '').strip()
            if not translation:
                logging.error("Kimi API 返回的译文为空。")
                return '翻译失败：Kimi API 返回的译文为空。'
            
            # 获取助手的回复并添加到消息历史中
            assistant_message = result.get('choices', [{}])[0].get('message', {})
            self.messages.append(assistant_message)
            
            logging.info(f"Kimi 翻译结果: {translation}")
            return translation
        except RateLimitError as e:
            logging.warning(f"Kimi API 触发速率限制: {e}")
            raise  # 重新抛出异常以触发重试
        except HTTPError as e:
            logging.error(f"调用 Kimi API 出错: {e}")
            return f'翻译失败：{str(e)}'
        except Exception as e:
            logging.error(f"解析 Kimi API 响应时出错: {e}")
            return f'翻译失败：{str(e)}'


    def translate_text(self):
        try:
            source_text = self.input_text.toPlainText().strip()
            if not source_text:
                QtWidgets.QMessageBox.warning(self, '警告', '请输入文本。')
                return

            from_lang = self.from_lang.currentText()
            to_lang = self.to_lang.currentText()

            logging.info(f"翻译请求 - 源语言: {from_lang}, 目标语言: {to_lang}, 文本: {source_text}")
            from tenacity import RetryError
            if from_lang == '自动检测':
                try:
                    detected_lang = detect(source_text)
                    if detected_lang.startswith('zh'):
                        from_lang = '中文'
                    elif detected_lang == 'en':
                        from_lang = '英文'
                    elif detected_lang == 'hi':
                        from_lang = '印地语'
                    elif detected_lang == 'ar':
                        from_lang = '阿拉伯语'
                    else:
                        from_lang = '自动检测'
                    logging.info(f"自动检测语言: {detected_lang}, 设置源语言为: {from_lang}")
                except LangDetectException:
                    from_lang = '自动检测'
                    logging.warning("语言检测失败，使用自动检测。")

            # 应用术语替换
            modified_text, terms = self.apply_key_terms(source_text, to_lang)
            logging.debug(f"应用术语替换后文本: {modified_text}, 术语: {terms}")

            # 将文本按句子分割
            sentences = self.split_sentences(modified_text, from_lang)
            logging.debug(f"分割后的句子: {sentences}")

            # 获取用户选择的翻译引擎
            engine_choice = self.engine_select.currentText()
            if engine_choice == 'Claude API':
                engine = 'claude'
            elif engine_choice == '百度翻译API':
                engine = 'baidu'
            elif engine_choice == 'Kimi API':
                engine = 'kimi'
            elif engine_choice == 'OpenAI API':
                engine = 'openai'  # 添加对 OpenAI 的支持
            else:
                QtWidgets.QMessageBox.warning(self, '警告', '未知的翻译引擎选择。')
                return

            translated_sentences = []
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                logging.info(f"处理句子: {sentence}")

                # 尝试从句子翻译记忆库中查找相似翻译
                similar_translation, similarity = self.find_similar_sentence(sentence)
                logging.debug(f"相似翻译: {similar_translation}, 相似度: {similarity}")

                if similar_translation:
                    translated_sentence = similar_translation
                    logging.info("使用记忆库中的翻译。")
                else:
                    if engine == 'claude':
                        # 使用 Claude API 进行翻译
                        machine_translation = self.claude_translate(sentence, from_lang, to_lang)
                        logging.info(f"Claude 翻译结果: {machine_translation}")
                    elif engine == 'baidu':
                        # 使用百度翻译API进行翻译
                        machine_translation = self.baidu_translate(sentence, from_lang, to_lang)
                        logging.info(f"Baidu 翻译结果: {machine_translation}")
                    elif engine == 'kimi':
                        # 使用 Kimi API 进行翻译
                        machine_translation = self.kimi_translate(sentence, from_lang, to_lang)
                        logging.info(f"Kimi 翻译结果: {machine_translation}")
                    elif engine == 'openai':
                        # 使用 OpenAI API 进行翻译
                        machine_translation = self.openai_translate(sentence, from_lang, to_lang)
                        logging.info(f"OpenAI 翻译结果: {machine_translation}")
                    else:
                        machine_translation = '翻译失败：未知的翻译引擎。'

                    if machine_translation.startswith('翻译失败') or machine_translation.startswith('请求翻译服务出错'):
                        translated_sentence = machine_translation
                    else:
                        # 进行词汇级别的替换
                        translated_sentence = self.replace_words(machine_translation, to_lang, from_lang)
                        logging.debug(f"词汇替换后译文: {translated_sentence}")

                        # 添加到翻译记忆库
                        self.add_sentence_translation(sentence, translated_sentence, from_lang, to_lang)
                        logging.info("将新翻译添加到翻译记忆库。")

                    translated_sentences.append((sentence, translated_sentence))

            # 合并译文
            if to_lang == '中文':
                final_translation = ''.join([t for s, t in translated_sentences])
            else:
                final_translation = ' '.join([t for s, t in translated_sentences])
            logging.debug(f"合并后的译文: {final_translation}")

            # 恢复术语
            final_translation = self.restore_key_terms(final_translation, terms)
            logging.debug(f"恢复术语后的最终译文: {final_translation}")

            # 限制 messages 列表的长度，保留最新的 20 条消息
            if len(self.messages) > 40:  # 每条对话包含用户和助手两条消息
                self.messages = self.messages[-40:]

            # 在输出区域显示源文本和译文，左右两栏
            self.output_text.clear()
            table_html = "<table border='1' width='100%'>"
            for source, target in translated_sentences:
                # 对HTML中的特殊字符进行转义
                source_escaped = source.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                target_escaped = target.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                table_html += f"<tr><td width='50%'>{source_escaped}</td><td width='50%'>{target_escaped}</td></tr>"
            table_html += "</table>"

            self.output_text.insertHtml(table_html)
            logging.info("译文显示完成。")
        except RetryError as re:
            QtWidgets.QMessageBox.critical(self, '错误', '翻译请求多次失败，可能由于API访问限制。请稍后再试或检查API配额。')
            logging.error(f"翻译过程中发生重试错误: {re}", exc_info=True)
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, '错误', f'翻译过程中发生错误：{str(e)}')
            logging.error(f"翻译过程中发生错误: {e}", exc_info=True)


    def apply_key_terms(self, text, to_lang='英文'):
        """
        根据术语库替换文本中的关键名词。
        """
        self.cursor.execute("SELECT source_term, target_term FROM key_terms WHERE to_lang=?", (to_lang,))
        terms = self.cursor.fetchall()
        if not terms:
            return text, []

        # 按长度排序，避免部分词汇被覆盖
        terms = sorted(terms, key=lambda x: len(x[0]), reverse=True)
        for source_term, target_term in terms:
            # 使用词边界匹配
            # 对中文，词边界可能不适用，直接替换nltk.download('punkt', quiet=True)
            # 对英文，使用 \b
            if to_lang == '英文':
                pattern = re.compile(rf'\b{re.escape(source_term)}\b')
                text = pattern.sub(f"__TERM_{source_term}__", text)
            else:
                text = text.replace(source_term, f"__TERM_{source_term}__")
        return text, terms

    def restore_key_terms(self, text, terms):
        """
        将占位符替换回术语的翻译。
        """
        for source_term, target_term in terms:
            placeholder = f"__TERM_{source_term}__"
            text = text.replace(placeholder, target_term)
        return text

    def split_sentences(self, text, lang):
        # 根据语言选择合适的分句方式
        if lang == '中文':
            # 使用中文的分句符号
            sentences = re.split('(?<=[。！？])', text)
        elif lang in ['印地语', '阿拉伯语']:
            # 使用印地语和阿拉伯语的分句符号
            sentences = re.split('(?<=[।؟!])', text)
        else:
            # 使用 sentence-splitter 进行句子分割
            try:
                lang_map = {
                    '英文': 'en',
                    '日文': 'ja',
                    '韩文': 'ko',
                    '法文': 'fr',
                    '德文': 'de',
                    '俄文': 'ru',
                    '西班牙文': 'es',
                    '印地语': 'hi',
                    '阿拉伯语': 'ar',
                }
                language_code = lang_map.get(lang, 'en')  # 默认为英文

                # 使用 sentence-splitter 进行分句
                sentences = split_text_into_sentences(text, language=language_code)
            except Exception as e:
                logging.warning(f"使用 sentence-splitter 分句时出错: {e}，将使用简单的正则表达式分句。")
                sentences = re.split(r'[.!?]', text)
        # 过滤掉空字符串
        sentences = [s.strip() for s in sentences if s.strip()]
        return sentences


    def display_side_by_side(self, sentence_pairs):
        # 清空输出区域
        self.output_text.clear()

        # 创建一个表格，左右两栏显示
        table_html = "<table border='1' width='100%'>"
        for source, target in sentence_pairs:
            # 对HTML中的特殊字符进行转义
            source_escaped = source.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            target_escaped = target.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            table_html += f"<tr><td width='50%'>{source_escaped}</td><td width='50%'>{target_escaped}</td></tr>"
        table_html += "</table>"

        self.output_text.insertHtml(table_html)

    def import_word_document(self):
        options = QtWidgets.QFileDialog.Options()
        file_path, _ = QtWidgets.QFileDialog.getOpenFileName(
            self,
            "选择 Word 文档",
            "",
            "Word Files (*.docx);;All Files (*)",
            options=options
        )
        if file_path:
            try:
                doc = Document(file_path)
                full_text = '\n'.join([para.text for para in doc.paragraphs])
                self.input_text.setText(full_text)
                QtWidgets.QMessageBox.information(self, '成功', 'Word 文档已导入。')
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, '错误', f'导入失败：{str(e)}')

    def export_to_word(self):
        options = QtWidgets.QFileDialog.Options()
        file_path, _ = QtWidgets.QFileDialog.getSaveFileName(
            self,
            "导出为 Word 文档",
            "",
            "Word Files (*.docx);;All Files (*)",
            options=options
        )
        if file_path:
            try:
                doc = Document()
                # 从输出区域获取HTML内容
                html_content = self.output_text.toHtml()

                # 将HTML内容转换为纯文本
                rows = re.findall(r'<tr>(.*?)</tr>', html_content, re.DOTALL)
                if not rows:
                    QtWidgets.QMessageBox.warning(self, '警告', '没有可导出的翻译结果。')
                    return
                table = doc.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '源文本'
                hdr_cells[1].text = '译文'
                for row in rows:
                    cols = re.findall(r'<td.*?>(.*?)</td>', row, re.DOTALL)
                    if len(cols) == 2:
                        source = re.sub('<.*?>', '', cols[0]).strip()
                        target = re.sub('<.*?>', '', cols[1]).strip()
                        row_cells = table.add_row().cells
                        row_cells[0].text = source
                        row_cells[1].text = target
                doc.save(file_path)
                QtWidgets.QMessageBox.information(self, '成功', '翻译结果已导出为 Word 文档。')
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, '错误', f'导出失败：{str(e)}')

    def import_memory(self):
        options = QtWidgets.QFileDialog.Options()
        file_path, _ = QtWidgets.QFileDialog.getOpenFileName(
            self,
            "选择翻译记忆库文件",
            "",
            "CSV Files (*.csv);;Excel Files (*.xlsx);;TMX Files (*.tmx);;Word Files (*.docx);;All Files (*)",
            options=options
        )
        if file_path:
            try:
                if file_path.lower().endswith('.csv'):
                    self.import_csv(file_path)
                elif file_path.lower().endswith('.xlsx'):
                    self.import_excel(file_path)
                elif file_path.lower().endswith('.tmx'):
                    self.import_tmx(file_path)
                elif file_path.lower().endswith('.docx'):
                    self.import_word(file_path)
                else:
                    QtWidgets.QMessageBox.warning(self, '警告', '不支持的文件格式。')
                    return
                QtWidgets.QMessageBox.information(self, '成功', '翻译记忆库已导入。')
                # 重新构建FAISS索引
                self.build_faiss_index()
                self.load_terminology_into_jieba()  # 重新加载术语库到 Jieba
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, '错误', f'导入失败：{str(e)}')

    def import_csv(self, file_path):
        df = pd.read_csv(file_path, encoding='utf-8')
        self.process_dataframe(df)

    def import_excel(self, file_path):
        df = pd.read_excel(file_path)
        self.process_dataframe(df)

    def import_tmx(self, file_path):
        from xml.etree import ElementTree as ET
        tree = ET.parse(file_path)
        root = tree.getroot()
        namespace = {'xml': 'http://www.w3.org/XML/1998/namespace'}
        tus = root.findall('.//tu')
        data = {'source': [], 'target': [], 'from_lang': [], 'to_lang': []}
        for tu in tus:
            tuv_elements = tu.findall('tuv')
            if len(tuv_elements) >= 2:
                source_tuv = tuv_elements[0]
                target_tuv = tuv_elements[1]
                source_lang = source_tuv.attrib.get('{http://www.w3.org/XML/1998/namespace}lang')
                target_lang = target_tuv.attrib.get('{http://www.w3.org/XML/1998/namespace}lang')
                source_text = source_tuv.find('seg').text
                target_text = target_tuv.find('seg').text
                data['source'].append(source_text)
                data['target'].append(target_text)
                data['from_lang'].append(source_lang)
                data['to_lang'].append(target_lang)
        df = pd.DataFrame(data)
        self.process_dataframe(df)

    def import_word(self, file_path):
        doc = Document(file_path)
        # 提取所有非空段落
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        # 假设翻译是交替排列的：源文本, 译文
        data = {'source': [], 'target': [], 'from_lang': [], 'to_lang': []}
        i = 0
        while i < len(paragraphs) - 1:
            source_text = paragraphs[i]
            target_text = paragraphs[i + 1]
            # 判断语言
            if any('\u4e00' <= char <= '\u9fff' for char in source_text):
                from_lang = '中文'
                to_lang = '英文'
            else:
                from_lang = '英文'
                to_lang = '中文'
            data['source'].append(source_text)
            data['target'].append(target_text)
            data['from_lang'].append(from_lang)
            data['to_lang'].append(to_lang)
            i += 2  # 处理下一对
        df = pd.DataFrame(data)
        self.process_dataframe(df)

    def process_dataframe(self, df):
        # 自动检测列名
        column_mapping = {
            'source': None,
            'target': None,
            'from_lang': None,
            'to_lang': None
        }
        for col in df.columns:
            lower_col = col.lower()
            if 'source' in lower_col or 'chinese' in lower_col or '源' in lower_col:
                column_mapping['source'] = col
            elif 'target' in lower_col or 'english' in lower_col or '目标' in lower_col:
                column_mapping['target'] = col
            elif 'from' in lower_col or '源语言' in lower_col:
                column_mapping['from_lang'] = col
            elif 'to' in lower_col or '目标语言' in lower_col:
                column_mapping['to_lang'] = col

        if column_mapping['source'] is None or column_mapping['target'] is None:
            raise ValueError('无法找到源文本或译文列。')

        if column_mapping['from_lang'] is None:
            df['from_lang'] = '自动检测'
            column_mapping['from_lang'] = 'from_lang'

        if column_mapping['to_lang'] is None:
            df['to_lang'] = '英文'
            column_mapping['to_lang'] = 'to_lang'

        for _, row in df.iterrows():
            source_text = str(row[column_mapping['source']]).strip()
            target_text = str(row[column_mapping['target']]).strip()
            from_lang = str(row[column_mapping['from_lang']]).strip()
            to_lang = str(row[column_mapping['to_lang']]).strip()
            if source_text and target_text:
                self.add_sentence_translation(source_text, target_text, from_lang, to_lang)

    def export_memory(self):
        options = QtWidgets.QFileDialog.Options()
        file_path, _ = QtWidgets.QFileDialog.getSaveFileName(
            self,
            "导出翻译记忆库",
            "",
            "CSV Files (*.csv);;Excel Files (*.xlsx);;TMX Files (*.tmx);;All Files (*)",
            options=options
        )
        if file_path:
            try:
                self.cursor.execute("SELECT source, target, from_lang, to_lang FROM sentence_translations")
                rows = self.cursor.fetchall()
                df = pd.DataFrame(rows, columns=['source', 'target', 'from_lang', 'to_lang'])
                if file_path.lower().endswith('.csv'):
                    df.to_csv(file_path, index=False, encoding='utf-8')
                elif file_path.lower().endswith('.xlsx'):
                    df.to_excel(file_path, index=False)
                elif file_path.lower().endswith('.tmx'):
                    self.export_tmx(df, file_path)
                else:
                    QtWidgets.QMessageBox.warning(self, '警告', '不支持的文件格式。')
                    return
                QtWidgets.QMessageBox.information(self, '成功', '翻译记忆库已导出。')
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, '错误', f'导出失败：{str(e)}')

    def export_tmx(self, df, file_path):
        from xml.etree.ElementTree import Element, SubElement, ElementTree
        root = Element('tmx', {'version': '1.4'})
        header = SubElement(root, 'header', {
            'creationtool': 'TranslatorApp',
            'creationtoolversion': '1.0',
            'segtype': 'sentence',
            'adminlang': 'en-us',
            'srclang': 'en',
            'datatype': 'PlainText'
        })
        body = SubElement(root, 'body')
        for _, row in df.iterrows():
            tu = SubElement(body, 'tu')
            tuv_source = SubElement(tu, 'tuv', {'xml:lang': row['from_lang']})
            seg_source = SubElement(tuv_source, 'seg')
            seg_source.text = row['source']
            tuv_target = SubElement(tu, 'tuv', {'xml:lang': row['to_lang']})
            seg_target = SubElement(tuv_target, 'seg')
            seg_target.text = row['target']
        tree = ElementTree(root)
        tree.write(file_path, encoding='utf-8', xml_declaration=True)

    def view_memory(self):
        # 创建并显示翻译记忆库窗口
        self.memory_window = MemoryWindow(self.conn)
        self.memory_window.show()

    def closeEvent(self, event):
        # 保存FAISS索引和ID映射
        faiss.write_index(self.index, self.faiss_index_file)
        with open(self.id_map_file, 'wb') as f:
            pickle.dump(self.id_map, f)
        self.conn.close()
        event.accept()

class MemoryWindow(QtWidgets.QWidget):
    def __init__(self, conn):
        super().__init__()
        self.conn = conn
        self.cursor = self.conn.cursor()
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle('翻译记忆库')
        self.resize(1000, 600)

        # Tabs for sentence, word translations, and key terms
        self.tabs = QtWidgets.QTabWidget()

        self.sentence_tab = QtWidgets.QWidget()
        self.word_tab = QtWidgets.QWidget()
        self.key_terms_tab = QtWidgets.QWidget()

        self.tabs.addTab(self.sentence_tab, '句子翻译')
        self.tabs.addTab(self.word_tab, '词汇翻译')
        self.tabs.addTab(self.key_terms_tab, '术语库')  # 新增术语库标签

        # 句子翻译表格
        self.init_sentence_tab()

        # 词汇翻译表格
        self.init_word_tab()

        # 术语库表格
        self.init_key_terms_tab()

        # 主布局
        main_layout = QtWidgets.QVBoxLayout()
        main_layout.addWidget(self.tabs)
        self.setLayout(main_layout)

    def init_sentence_tab(self):
        self.sentence_layout = QtWidgets.QVBoxLayout()
        self.sentence_table = QtWidgets.QTableWidget()
        self.sentence_table.setColumnCount(5)
        self.sentence_table.setHorizontalHeaderLabels(['ID', '源文本', '译文', '源语言', '目标语言'])
        self.sentence_layout.addWidget(self.sentence_table)

        self.sentence_button_layout = QtWidgets.QHBoxLayout()
        self.edit_sentence_button = QtWidgets.QPushButton('编辑')
        self.delete_sentence_button = QtWidgets.QPushButton('删除')
        self.sentence_button_layout.addWidget(self.edit_sentence_button)
        self.sentence_button_layout.addWidget(self.delete_sentence_button)
        self.sentence_layout.addLayout(self.sentence_button_layout)

        self.sentence_tab.setLayout(self.sentence_layout)

        # 加载数据
        self.load_sentence_translations()

        # 连接按钮
        self.edit_sentence_button.clicked.connect(self.edit_sentence)
        self.delete_sentence_button.clicked.connect(self.delete_sentence)

    def init_word_tab(self):
        self.word_layout = QtWidgets.QVBoxLayout()
        self.word_table = QtWidgets.QTableWidget()
        self.word_table.setColumnCount(5)
        self.word_table.setHorizontalHeaderLabels(['ID', '源词', '译词', '源语言', '目标语言'])
        self.word_layout.addWidget(self.word_table)

        self.word_button_layout = QtWidgets.QHBoxLayout()
        self.edit_word_button = QtWidgets.QPushButton('编辑')
        self.delete_word_button = QtWidgets.QPushButton('删除')
        self.word_button_layout.addWidget(self.edit_word_button)
        self.word_button_layout.addWidget(self.delete_word_button)
        self.word_layout.addLayout(self.word_button_layout)

        self.word_tab.setLayout(self.word_layout)

        # 加载数据
        self.load_word_translations()

        # 连接按钮
        self.edit_word_button.clicked.connect(self.edit_word)
        self.delete_word_button.clicked.connect(self.delete_word)

    def init_key_terms_tab(self):
        self.key_terms_layout = QtWidgets.QVBoxLayout()
        self.key_terms_table = QtWidgets.QTableWidget()
        self.key_terms_table.setColumnCount(5)
        self.key_terms_table.setHorizontalHeaderLabels(['ID', '源术语', '译术语', '源语言', '目标语言'])
        self.key_terms_layout.addWidget(self.key_terms_table)

        self.key_terms_button_layout = QtWidgets.QHBoxLayout()
        self.edit_key_term_button = QtWidgets.QPushButton('编辑')
        self.delete_key_term_button = QtWidgets.QPushButton('删除')
        self.key_terms_button_layout.addWidget(self.edit_key_term_button)
        self.key_terms_button_layout.addWidget(self.delete_key_term_button)
        self.key_terms_layout.addLayout(self.key_terms_button_layout)

        self.key_terms_tab.setLayout(self.key_terms_layout)

        # 加载数据
        self.load_key_terms()

        # 连接按钮
        self.edit_key_term_button.clicked.connect(self.edit_key_term)
        self.delete_key_term_button.clicked.connect(self.delete_key_term)

    def load_sentence_translations(self):
        self.cursor.execute("SELECT id, source, target, from_lang, to_lang FROM sentence_translations")
        rows = self.cursor.fetchall()
        self.sentence_table.setRowCount(len(rows))
        for row_idx, row in enumerate(rows):
            for col_idx, item in enumerate(row):
                table_item = QtWidgets.QTableWidgetItem(str(item))
                table_item.setFlags(table_item.flags() ^ QtCore.Qt.ItemIsEditable)
                self.sentence_table.setItem(row_idx, col_idx, table_item)
        self.sentence_table.resizeColumnsToContents()

    def load_word_translations(self):
        self.cursor.execute("SELECT id, source, target, from_lang, to_lang FROM word_translations")
        rows = self.cursor.fetchall()
        self.word_table.setRowCount(len(rows))
        for row_idx, row in enumerate(rows):
            for col_idx, item in enumerate(row):
                table_item = QtWidgets.QTableWidgetItem(str(item))
                table_item.setFlags(table_item.flags() ^ QtCore.Qt.ItemIsEditable)
                self.word_table.setItem(row_idx, col_idx, table_item)
        self.word_table.resizeColumnsToContents()

    def load_key_terms(self):
        self.cursor.execute("SELECT id, source_term, target_term, from_lang, to_lang FROM key_terms")
        rows = self.cursor.fetchall()
        self.key_terms_table.setRowCount(len(rows))
        for row_idx, row in enumerate(rows):
            for col_idx, item in enumerate(row):
                table_item = QtWidgets.QTableWidgetItem(str(item))
                table_item.setFlags(table_item.flags() ^ QtCore.Qt.ItemIsEditable)
                self.key_terms_table.setItem(row_idx, col_idx, table_item)
        self.key_terms_table.resizeColumnsToContents()

    def edit_sentence(self):
        selected_items = self.sentence_table.selectedItems()
        if not selected_items:
            QtWidgets.QMessageBox.warning(self, '警告', '请选择要编辑的句子。')
            return
        row = selected_items[0].row()
        id_ = int(self.sentence_table.item(row, 0).text())
        source = self.sentence_table.item(row, 1).text()
        target = self.sentence_table.item(row, 2).text()
        from_lang = self.sentence_table.item(row, 3).text()
        to_lang = self.sentence_table.item(row, 4).text()

        dialog = EditSentenceDialog(id_, source, target, from_lang, to_lang, self.conn)
        if dialog.exec_():
            self.load_sentence_translations()

    def delete_sentence(self):
        selected_items = self.sentence_table.selectedItems()
        if not selected_items:
            QtWidgets.QMessageBox.warning(self, '警告', '请选择要删除的句子。')
            return
        row = selected_items[0].row()
        id_ = int(self.sentence_table.item(row, 0).text())
        reply = QtWidgets.QMessageBox.question(
            self, '确认删除', '确定要删除选中的句子翻译吗？',
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No, QtWidgets.QMessageBox.No
        )
        if reply == QtWidgets.QMessageBox.Yes:
            self.cursor.execute("DELETE FROM sentence_translations WHERE id=?", (id_,))
            self.conn.commit()
            self.load_sentence_translations()

    def edit_word(self):
        selected_items = self.word_table.selectedItems()
        if not selected_items:
            QtWidgets.QMessageBox.warning(self, '警告', '请选择要编辑的词汇。')
            return
        row = selected_items[0].row()
        id_ = int(self.word_table.item(row, 0).text())
        source = self.word_table.item(row, 1).text()
        target = self.word_table.item(row, 2).text()
        from_lang = self.word_table.item(row, 3).text()
        to_lang = self.word_table.item(row, 4).text()

        dialog = EditWordDialog(id_, source, target, from_lang, to_lang, self.conn)
        if dialog.exec_():
            self.load_word_translations()

    def delete_word(self):
        selected_items = self.word_table.selectedItems()
        if not selected_items:
            QtWidgets.QMessageBox.warning(self, '警告', '请选择要删除的词汇。')
            return
        row = selected_items[0].row()
        id_ = int(self.word_table.item(row, 0).text())
        reply = QtWidgets.QMessageBox.question(
            self, '确认删除', '确定要删除选中的词汇翻译吗？',
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No, QtWidgets.QMessageBox.No
        )
        if reply == QtWidgets.QMessageBox.Yes:
            self.cursor.execute("DELETE FROM word_translations WHERE id=?", (id_,))
            self.conn.commit()
            self.load_word_translations()

    def edit_key_term(self):
        selected_items = self.key_terms_table.selectedItems()
        if not selected_items:
            QtWidgets.QMessageBox.warning(self, '警告', '请选择要编辑的术语。')
            return
        row = selected_items[0].row()
        id_ = int(self.key_terms_table.item(row, 0).text())
        source_term = self.key_terms_table.item(row, 1).text()
        target_term = self.key_terms_table.item(row, 2).text()
        from_lang = self.key_terms_table.item(row, 3).text()
        to_lang = self.key_terms_table.item(row, 4).text()

        dialog = EditKeyTermDialog(id_, source_term, target_term, from_lang, to_lang, self.conn)
        if dialog.exec_():
            self.load_key_terms()

    def delete_key_term(self):
        selected_items = self.key_terms_table.selectedItems()
        if not selected_items:
            QtWidgets.QMessageBox.warning(self, '警告', '请选择要删除的术语。')
            return
        row = selected_items[0].row()
        id_ = int(self.key_terms_table.item(row, 0).text())
        reply = QtWidgets.QMessageBox.question(
            self, '确认删除', '确定要删除选中的术语吗？',
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No, QtWidgets.QMessageBox.No
        )
        if reply == QtWidgets.QMessageBox.Yes:
            self.cursor.execute("DELETE FROM key_terms WHERE id=?", (id_,))
            self.conn.commit()
            self.load_key_terms()

class EditSentenceDialog(QtWidgets.QDialog):
    def __init__(self, id_, source, target, from_lang, to_lang, conn):
        super().__init__()
        self.id_ = id_
        self.conn = conn
        self.cursor = self.conn.cursor()
        self.setWindowTitle('编辑句子翻译')
        self.init_ui(source, target, from_lang, to_lang)

    def init_ui(self, source, target, from_lang, to_lang):
        layout = QtWidgets.QFormLayout()

        self.source_edit = QtWidgets.QTextEdit(source)
        self.target_edit = QtWidgets.QTextEdit(target)
        self.from_lang_edit = QtWidgets.QComboBox()
        self.from_lang_edit.addItems(['自动检测', '中文', '英文'])
        self.from_lang_edit.setCurrentText(from_lang)

        self.to_lang_edit = QtWidgets.QComboBox()
        self.to_lang_edit.addItems(['中文', '英文'])
        self.to_lang_edit.setCurrentText(to_lang)

        layout.addRow('源文本：', self.source_edit)
        layout.addRow('译文：', self.target_edit)
        layout.addRow('源语言：', self.from_lang_edit)
        layout.addRow('目标语言：', self.to_lang_edit)

        self.save_button = QtWidgets.QPushButton('保存')
        self.save_button.clicked.connect(self.save_changes)
        self.cancel_button = QtWidgets.QPushButton('取消')
        self.cancel_button.clicked.connect(self.reject)

        button_layout = QtWidgets.QHBoxLayout()
        button_layout.addWidget(self.save_button)
        button_layout.addWidget(self.cancel_button)
        layout.addRow(button_layout)

        self.setLayout(layout)

    def save_changes(self):
        source = self.source_edit.toPlainText().strip()
        target = self.target_edit.toPlainText().strip()
        from_lang = self.from_lang_edit.currentText()
        to_lang = self.to_lang_edit.currentText()

        if not source or not target:
            QtWidgets.QMessageBox.warning(self, '警告', '源文本和译文不能为空。')
            return

        self.cursor.execute('''
            UPDATE sentence_translations
            SET source=?, target=?, from_lang=?, to_lang=?
            WHERE id=?
        ''', (source, target, from_lang, to_lang, self.id_))
        self.conn.commit()
        self.accept()

class EditWordDialog(QtWidgets.QDialog):
    def __init__(self, id_, source, target, from_lang, to_lang, conn):
        super().__init__()
        self.id_ = id_
        self.conn = conn
        self.cursor = self.conn.cursor()
        self.setWindowTitle('编辑词汇翻译')
        self.init_ui(source, target, from_lang, to_lang)

    def init_ui(self, source, target, from_lang, to_lang):
        layout = QtWidgets.QFormLayout()

        self.source_edit = QtWidgets.QLineEdit(source)
        self.target_edit = QtWidgets.QLineEdit(target)
        self.from_lang_edit = QtWidgets.QComboBox()
        self.from_lang_edit.addItems(['中文', '英文'])
        self.from_lang_edit.setCurrentText(from_lang)

        self.to_lang_edit = QtWidgets.QComboBox()
        self.to_lang_edit.addItems(['中文', '英文'])
        self.to_lang_edit.setCurrentText(to_lang)

        layout.addRow('源词：', self.source_edit)
        layout.addRow('译词：', self.target_edit)
        layout.addRow('源语言：', self.from_lang_edit)
        layout.addRow('目标语言：', self.to_lang_edit)

        self.save_button = QtWidgets.QPushButton('保存')
        self.save_button.clicked.connect(self.save_changes)
        self.cancel_button = QtWidgets.QPushButton('取消')
        self.cancel_button.clicked.connect(self.reject)

        button_layout = QtWidgets.QHBoxLayout()
        button_layout.addWidget(self.save_button)
        button_layout.addWidget(self.cancel_button)
        layout.addRow(button_layout)

        self.setLayout(layout)

    def save_changes(self):
        source = self.source_edit.text().strip()
        target = self.target_edit.text().strip()
        from_lang = self.from_lang_edit.currentText()
        to_lang = self.to_lang_edit.currentText()

        if not source or not target:
            QtWidgets.QMessageBox.warning(self, '警告', '源词和译词不能为空。')
            return

        self.cursor.execute('''
            UPDATE word_translations
            SET source=?, target=?, from_lang=?, to_lang=?
            WHERE id=?
        ''', (source, target, from_lang, to_lang, self.id_))
        self.conn.commit()
        self.accept()

class EditKeyTermDialog(QtWidgets.QDialog):
    def __init__(self, id_, source_term, target_term, from_lang, to_lang, conn):
        super().__init__()
        self.id_ = id_
        self.conn = conn
        self.cursor = self.conn.cursor()
        self.setWindowTitle('编辑术语')
        self.init_ui(source_term, target_term, from_lang, to_lang)

    def init_ui(self, source_term, target_term, from_lang, to_lang):
        layout = QtWidgets.QFormLayout()

        self.source_edit = QtWidgets.QLineEdit(source_term)
        self.target_edit = QtWidgets.QLineEdit(target_term)
        self.from_lang_edit = QtWidgets.QComboBox()
        self.from_lang_edit.addItems(['中文', '英文'])
        self.from_lang_edit.setCurrentText(from_lang)

        self.to_lang_edit = QtWidgets.QComboBox()
        self.to_lang_edit.addItems(['中文', '英文'])
        self.to_lang_edit.setCurrentText(to_lang)

        layout.addRow('源术语：', self.source_edit)
        layout.addRow('译术语：', self.target_edit)
        layout.addRow('源语言：', self.from_lang_edit)
        layout.addRow('目标语言：', self.to_lang_edit)

        self.save_button = QtWidgets.QPushButton('保存')
        self.save_button.clicked.connect(self.save_changes)
        self.cancel_button = QtWidgets.QPushButton('取消')
        self.cancel_button.clicked.connect(self.reject)

        button_layout = QtWidgets.QHBoxLayout()
        button_layout.addWidget(self.save_button)
        button_layout.addWidget(self.cancel_button)
        layout.addRow(button_layout)

        self.setLayout(layout)

    def save_changes(self):
        source_term = self.source_edit.text().strip()
        target_term = self.target_edit.text().strip()
        from_lang = self.from_lang_edit.currentText()
        to_lang = self.to_lang_edit.currentText()

        if not source_term or not target_term:
            QtWidgets.QMessageBox.warning(self, '警告', '源术语和译术语不能为空。')
            return

        try:
            self.cursor.execute('''
                UPDATE key_terms
                SET source_term=?, target_term=?, from_lang=?, to_lang=?
                WHERE id=?
            ''', (source_term, target_term, from_lang, to_lang, self.id_))
            self.conn.commit()
            self.accept()
        except sqlite3.IntegrityError:
            QtWidgets.QMessageBox.warning(self, '警告', '源术语已存在。')

    
if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    translator = TranslatorApp()  # 确保这个 TranslatorApp 是您定义的类
    sys.exit(app.exec_())